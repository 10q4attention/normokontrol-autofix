"""
Фиксер 06: структурные правки документа.
- Преобразует «Содержание»/«Оглавление» в заголовок 1 «ОГЛАВЛЕНИЕ»
- Корректирует уровни ВСЕХ заголовков по тексту:
    именованные главные разделы (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ …)  → Heading 1
    X.   НАЗВАНИЕ (нумерованная глава)                     → Heading 1
    X.Y  название                                          → Heading 2
    X.Y.Z название                                         → Heading 3
  + применяет форматирование нового уровня
- Исправляет ПРОПИСНЫЕ в TOC-параграфах (MAIN_IN_TOC)
"""

import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from fixers.base_fixer import BaseFixer, FixResult

# Паттерны для определения уровня по тексту
_L3_RE = re.compile(r'^\d+\.\d+\.\d+[\.\s]')   # 1.2.3. или 1.2.3 текст
_L2_RE = re.compile(r'^\d+\.\d+[\.\s]')          # 1.2. или 1.2 текст
_L1_RE = re.compile(r'^\d+\.\s')                  # 1. НАЗВАНИЕ

# Именованные разделы, которые всегда Heading 1
_MAIN_H1 = frozenset({
    'введение', 'заключение', 'аннотация', 'реферат', 'приложения',
    'список использованных источников', 'список литературы',
    'список используемых сокращений и обозначений',
})

# Разделы в TOC, которые должны быть ПРОПИСНЫМИ
_MAIN_IN_TOC = [
    'ВВЕДЕНИЕ', 'ЗАКЛЮЧЕНИЕ', 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ',
    'СПИСОК ЛИТЕРАТУРЫ', 'ПРИЛОЖЕНИЯ', 'АННОТАЦИЯ', 'РЕФЕРАТ',
    'ТЕОРЕТИЧЕСКИЙ РАЗДЕЛ', 'ПРАКТИЧЕСКИЙ РАЗДЕЛ', 'ТЕХНОЛОГИЧЕСКИЙ РАЗДЕЛ',
]

_H1_STYLES = ('heading 1', 'заголовок 1')
_H2_STYLES = ('heading 2', 'заголовок 2')
_H3_STYLES = ('heading 3', 'заголовок 3')
_H4_STYLES = ('heading 4', 'заголовок 4')

# Форматирование по уровню (те же значения, что в fixer_05 _EXP)
_LEVEL_FMT = {
    1: {'size': 18, 'sb': 0,  'sa': 12, 'indent': 1.25},
    2: {'size': 16, 'sb': 24, 'sa': 12, 'indent': 1.25},
    3: {'size': 14, 'sb': 24, 'sa': 12, 'indent': 1.25},
}


def _get_style(doc, candidates):
    for name in doc.styles:
        if name.name.lower() in candidates:
            return name
    return None


def _current_level(para):
    style = (para.style.name or '').lower() if para.style else ''
    for lvl, names in [(1, _H1_STYLES), (2, _H2_STYLES), (3, _H3_STYLES), (4, _H4_STYLES)]:
        if any(s in style for s in names):
            return lvl
    return None


def _target_level(text: str):
    """Определяет нужный уровень заголовка по его тексту. None = неизвестно."""
    t = text.strip()
    if t.lower() in _MAIN_H1:
        return 1
    if _L3_RE.match(t):
        return 3
    if _L2_RE.match(t):
        return 2
    if _L1_RE.match(t):
        return 1
    return None


def _style_for_level(doc, level: int):
    mapping = {1: _H1_STYLES, 2: _H2_STYLES, 3: _H3_STYLES}
    return _get_style(doc, mapping.get(level, _H3_STYLES))


def _apply_heading_fmt(para, level: int, is_main: bool = False):
    """Применяет форматирование согласно уровню."""
    lvl = min(max(level, 1), 3)
    fmt = _LEVEL_FMT[lvl]
    pf = para.paragraph_format

    if is_main:
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
    else:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(fmt['indent'])

    pf.left_indent = Pt(0)
    pf.space_before = Pt(fmt['sb'])
    pf.space_after = Pt(fmt['sa'])
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5

    size = Pt(fmt['size'])
    for run in para.runs:
        if not run.text:
            continue
        run.font.name = 'Times New Roman'
        run.font.size = size
        run.font.bold = True
        run.font.italic = False
        if is_main or level == 1:
            run.text = run.text.upper()


class StructureFixer(BaseFixer):
    @property
    def fixer_id(self): return "structure"

    @property
    def name(self): return "Структура документа"

    def fix(self, doc, model) -> FixResult:
        result = FixResult(fixer_id=self.fixer_id, name=self.name)

        # ── 1. ОГЛАВЛЕНИЕ ────────────────────────────────────────────
        if self._fix_ogl_heading(doc):
            result.changes.append("Заголовок «ОГЛАВЛЕНИЕ» исправлен/добавлен")

        # ── 2. Уровни заголовков ─────────────────────────────────────
        counts = {1: 0, 2: 0, 3: 0}
        for para in doc.paragraphs:
            cur = _current_level(para)
            if cur is None:
                continue
            text = para.text.strip()
            if not text:
                continue
            # ОГЛАВЛЕНИЕ уже обработано выше
            if text.upper() in ('ОГЛАВЛЕНИЕ', 'СОДЕРЖАНИЕ'):
                continue

            tgt = _target_level(text)
            if tgt is None or tgt == cur:
                continue

            new_style = _style_for_level(doc, tgt)
            if new_style is None:
                continue

            para.style = new_style
            is_main = text.lower() in _MAIN_H1
            _apply_heading_fmt(para, tgt, is_main=is_main)
            counts[tgt] += 1

        for lvl, cnt in counts.items():
            if cnt:
                result.changes.append(f"Заголовки переведены на уровень {lvl}: {cnt}")

        # ── 3. ПРОПИСНЫЕ в TOC-параграфах ────────────────────────────
        toc_fixed = self._fix_toc_uppercase(doc)
        if toc_fixed:
            result.changes.append(f"ПРОПИСНЫЕ в оглавлении исправлены: {toc_fixed}")

        if not result.changes:
            result.status = 'skipped'
        return result

    def _fix_ogl_heading(self, doc) -> bool:
        h1_style = _get_style(doc, _H1_STYLES)
        if h1_style is None:
            return False

        for para in doc.paragraphs:
            text = para.text.strip().lower()
            if text not in ('содержание', 'оглавление'):
                continue

            para.style = h1_style

            found_first = False
            for run in para.runs:
                if run.text.strip() and not found_first:
                    run.text = 'ОГЛАВЛЕНИЕ'
                    found_first = True
                elif run.text.strip():
                    run.text = ''

            if not found_first:
                from docx.oxml import OxmlElement
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = 'ОГЛАВЛЕНИЕ'
                r.append(t)
                para._element.append(r)

            pf = para.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.space_before = Pt(0)
            pf.space_after = Pt(12)
            pf.first_line_indent = Cm(0)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.5

            for run in para.runs:
                if run.text.strip():
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.italic = False

            return True

        return False

    def _fix_toc_uppercase(self, doc) -> int:
        """В параграфах стиля TOC заменяет названия главных разделов на ПРОПИСНЫЕ."""
        fixed = 0
        _toc_styles = ('toc', 'оглавление', 'содержание', 'contents')

        for para in doc.paragraphs:
            style = (para.style.name or '').lower() if para.style else ''
            # TOC-параграфы — либо стиль TOC, либо строки типа "ВВЕДЕНИЕ\t3"
            is_toc_style = any(s in style for s in _toc_styles)
            is_toc_line = bool(re.search(r'(\t\s*\d+\s*$|\.{2,}\s*\d+\s*$)', para.text))
            if not (is_toc_style or is_toc_line):
                continue

            text = para.text
            changed = False
            for section in _MAIN_IN_TOC:
                idx = text.upper().find(section)
                if idx < 0:
                    continue
                actual = text[idx:idx + len(section)]
                if actual != section:
                    text = text[:idx] + section + text[idx + len(section):]
                    changed = True

            if not changed:
                continue

            # Применяем исправленный текст к runs
            full = para.text
            if full == text:
                continue
            # Перераспределяем текст: всё в первый непустой run
            runs = [r for r in para.runs if r.text]
            if not runs:
                continue
            runs[0].text = text
            for r in runs[1:]:
                r.text = ''
            fixed += 1

        return fixed
