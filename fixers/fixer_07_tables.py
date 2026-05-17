"""Фиксер 07: подписи к таблицам + форматирование ячеек."""

import re
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from fixers.base_fixer import BaseFixer, FixResult

_TABLE_CAP_RE = re.compile(
    r'^(?:Таблица|Продолжение\s+[Тт]аблицы)\s+[\dА-ЯA-Z]+\.\d+',
    re.UNICODE
)


def _fix_table_caption(para):
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.0

    for run in para.runs:
        if not run.text:
            continue
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.bold = False
        t = run.text
        if t.strip() == '-':
            run.text = t.replace('-', '—')
        elif ' - ' in t:
            run.text = t.replace(' - ', ' — ')


def _fix_table_cells(table):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                pf = para.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.first_line_indent = Pt(0)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.line_spacing = 1.0

                if ri == 0:
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif ci == 0:
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                for run in para.runs:
                    if not run.text:
                        continue
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)


class TableCaptionFixer(BaseFixer):
    @property
    def fixer_id(self): return "table_captions"

    @property
    def name(self): return "Подписи и ячейки таблиц"

    def fix(self, doc, model) -> FixResult:
        result = FixResult(fixer_id=self.fixer_id, name=self.name)

        # Подписи
        fixed_caps = 0
        for para in doc.paragraphs:
            if _TABLE_CAP_RE.match(para.text.strip()):
                _fix_table_caption(para)
                fixed_caps += 1

        # Ячейки таблиц у которых есть подписи — матчинг по XML-элементу,
        # т.к. doc.tables включает вложенные таблицы и индексы не совпадают
        captioned_tbls = {t['_tbl'] for t in model.tables if t.get('linked_from') and t.get('_tbl') is not None}
        fixed_tables = 0
        for table in doc.tables:
            if table._tbl in captioned_tbls:
                _fix_table_cells(table)
                fixed_tables += 1

        # п.7.5: абзац, следующий за таблицей, должен иметь space_before=6
        fixed_after = 0
        paras = doc.paragraphs
        para_map = {p._element: p for p in paras}
        body = doc.element.body
        children = list(body)
        for idx, child in enumerate(children):
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'tbl':
                # Найти следующий параграф в body
                for next_child in children[idx+1:]:
                    ntag = next_child.tag.split('}')[-1] if '}' in next_child.tag else next_child.tag
                    if ntag == 'p':
                        np = para_map.get(next_child)
                        if np and (np.paragraph_format.space_before is None or
                                   np.paragraph_format.space_before.pt < 5):
                            np.paragraph_format.space_before = Pt(6)
                            fixed_after += 1
                        break

        if fixed_caps:
            result.changes.append(f"Исправлено подписей таблиц: {fixed_caps}")
        if fixed_tables:
            result.changes.append(f"Исправлено форматирование ячеек в таблицах: {fixed_tables}")
        if fixed_after:
            result.changes.append(f"Установлен интервал перед абзацем после таблицы: {fixed_after}")

        if not result.changes:
            result.status = 'skipped'
        return result
